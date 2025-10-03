# Copyright Daniel Stancl.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

from __future__ import annotations

import copy
import csv
import tkinter as tk
from typing import TYPE_CHECKING

import pytest
from docx import Document

if TYPE_CHECKING:
    import pathlib
    from collections.abc import Iterator


@pytest.fixture
def boring_iterator() -> Iterator[list[str]]:
    return iter(
        [
            ["", "v1", "v2", "v3"],
            ["Field", "Name", "City", "Agreement Price"],
            ["", "Mr Adel", "Prague", "200.000"],
        ],
    )


@pytest.fixture
def boring_csv(tmp_path: pathlib.Path, boring_iterator: Iterator[list[str]]) -> str:
    csv_path = tmp_path / "boring.csv"
    with open(csv_path, "w") as f:
        writer = csv.writer(f)
        # Need to copy iterator not to exhaust it for tests
        for row in copy.copy(boring_iterator):
            writer.writerow(row)
    return str(csv_path)


@pytest.fixture
def boring_document(tmp_path: pathlib.Path) -> str:
    doc_path = tmp_path / "boring.docx"
    doc = Document()

    doc.add_paragraph("This is first paragraph {v1} and {v2}.")
    paragraph = doc.add_paragraph("This is second paragraph {v3}.")
    paragraph.add_run("bold").bold = True
    paragraph.add_run("italic").italic = True

    doc.save(doc_path)
    return str(doc_path)


@pytest.fixture
def root() -> Iterator[tk.Tk]:
    root = tk.Tk()
    # Hide the root window for testing
    root.withdraw()
    yield root

    root.destroy()

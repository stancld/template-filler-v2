# Copyright Daniel Stancl.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

from __future__ import annotations

import os
from typing import TYPE_CHECKING

import pytest
from docx import Document

from template_filler.data_loading import read_csv
from template_filler.word_filler import fill_template

if TYPE_CHECKING:
    import pathlib


@pytest.mark.unit
def test_fill_template(boring_document: str, boring_csv: str, tmp_path: pathlib.Path) -> None:
    output_path = tmp_path / "output.docx"
    fill_template(boring_document, str(output_path), next(read_csv(boring_csv)))
    assert os.path.exists(output_path)
    assert os.path.isfile(output_path)

    output_doc = Document(output_path)
    assert output_doc.paragraphs[0].text == "This is first paragraph Mr Adel and Prague."
    assert output_doc.paragraphs[1].text == "This is second paragraph 200.000.bolditalic"
    assert output_doc.paragraphs[1].runs[0].text == "This is second paragraph 200.000."
    assert output_doc.paragraphs[1].runs[1].text == "bold"
    assert output_doc.paragraphs[1].runs[1].bold
    assert output_doc.paragraphs[1].runs[2].text == "italic"
    assert output_doc.paragraphs[1].runs[2].italic


@pytest.mark.unit
def test_fill_template_partial_replacement(tmp_path: pathlib.Path) -> None:
    """Test that fill_template works with partial replacement."""
    doc_path = tmp_path / "partial.docx"
    doc = Document()
    doc.add_paragraph("Values: {v1}, {v2}, {v3}")
    doc.save(doc_path)
    output_path = tmp_path / "output_partial.docx"
    fill_template(str(doc_path), str(output_path), {"v1": "Value1", "v3": "Value3"})
    output_doc = Document(output_path)
    assert output_doc.paragraphs[0].text == "Values: Value1, {v2}, Value3"


@pytest.mark.unit
def test_fill_template_invalid_path(tmp_path: pathlib.Path) -> None:
    """Test that fill_template handles invalid paths gracefully."""
    output_path = tmp_path / "output_invalid.docx"
    with pytest.raises(Exception):
        fill_template("nonexistent_file.docx", str(output_path), {"v1": "Value"})
